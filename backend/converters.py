import io

import openpyxl


def _runs_to_md(runs) -> str:
    md = ""
    for run in runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        md += t
    return md


def _table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    header_cells = [c.text.strip().replace("\n", " ").replace("|", "\\|") for c in rows[0].cells]
    header = "| " + " | ".join(header_cells) + " |"
    sep = "| " + " | ".join("---" for _ in header_cells) + " |"
    data_rows = []
    for row in rows[1:]:
        cells = [c.text.strip().replace("\n", " ").replace("|", "\\|") for c in row.cells]
        data_rows.append("| " + " | ".join(cells) + " |")
    return "\n".join([header, sep] + data_rows)


HEADING_STYLES = {
    "Heading 1": "#", "Heading 2": "##", "Heading 3": "###",
    "Heading 4": "####", "Heading 5": "#####", "Heading 6": "######",
    "Title": "#", "Subtitle": "##",
}


def docx_to_markdown(content: bytes) -> tuple[str, dict]:
    from docx import Document

    doc = Document(io.BytesIO(content))
    parts = []
    info = {"type": "word", "paragraphs": 0, "headings": 0, "tables": 0}

    para_idx = 0
    table_idx = 0

    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if tag == "p":
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else ""
                if style_name in HEADING_STYLES:
                    parts.append(f"{HEADING_STYLES[style_name]} {text}")
                    info["headings"] += 1
                else:
                    md_text = _runs_to_md(para.runs)
                    if md_text.strip():
                        parts.append(md_text)
                        info["paragraphs"] += 1

        elif tag == "tbl":
            if table_idx < len(doc.tables):
                tbl = doc.tables[table_idx]
                table_idx += 1
                md_table = _table_to_md(tbl)
                if md_table:
                    parts.append(md_table)
                    info["tables"] += 1

    return "\n\n".join(parts), info


def xlsx_to_markdown(content: bytes) -> tuple[str, dict]:
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    parts = []
    info = {"type": "excel", "sheets": []}

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        non_empty = [r for r in rows if any(c is not None for c in r)]

        sheet_info = {
            "name": sheet_name,
            "rows": len(non_empty),
            "cols": len(non_empty[0]) if non_empty else 0,
        }
        info["sheets"].append(sheet_info)

        if not non_empty:
            continue

        parts.append(f"## Sheet: {sheet_name}")

        header = non_empty[0]
        parts.append("| " + " | ".join((str(c) if c is not None else "").replace("\n", " ").replace("|", "\\|") for c in header) + " |")
        parts.append("| " + " | ".join("---" for _ in header) + " |")

        for row in non_empty[1:]:
            parts.append("| " + " | ".join((str(c) if c is not None else "").replace("\n", " ").replace("|", "\\|") for c in row) + " |")

    return "\n\n".join(parts), info


def convert_to_pdf_bytes(content: bytes, filename: str) -> bytes:
    import subprocess
    import tempfile
    import os

    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, filename)
        with open(input_path, "wb") as f:
            f.write(content)

        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", tmpdir, input_path],
            capture_output=True,
            timeout=120,
        )

        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice conversion failed: {result.stderr.decode(errors='replace')}"
            )

        base = os.path.splitext(filename)[0]
        pdf_path = os.path.join(tmpdir, base + ".pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice did not produce a PDF file")

        with open(pdf_path, "rb") as f:
            return f.read()
