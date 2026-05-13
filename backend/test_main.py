import io
import shutil

import pytest
from docx import Document
from fastapi.testclient import TestClient
import openpyxl

from main import app

client = TestClient(app)


def make_docx_bytes() -> bytes:
    doc = Document()
    doc.add_heading("Test Title", level=1)
    doc.add_paragraph("Hello world paragraph.")
    tbl = doc.add_table(rows=2, cols=2)
    tbl.rows[0].cells[0].text = "Col A"
    tbl.rows[0].cells[1].text = "Col B"
    tbl.rows[1].cells[0].text = "Val 1"
    tbl.rows[1].cells[1].text = "Val 2"
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def make_xlsx_bytes() -> bytes:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Sales"
    ws.append(["Name", "Amount"])
    ws.append(["Alice", 100])
    ws.append(["Bob", 200])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_health():
    r = client.get("/api/health")
    assert r.status_code == 200
    assert r.json() == {"status": "ok"}


def test_convert_direct_docx():
    docx_bytes = make_docx_bytes()
    r = client.post(
        "/api/convert-direct",
        files={"file": ("document.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
    )
    assert r.status_code == 200
    data = r.json()
    assert "markdown" in data
    assert "info" in data
    assert data["filename"] == "document.docx"
    assert "# Test Title" in data["markdown"]
    assert "Hello world paragraph" in data["markdown"]
    assert "Col A" in data["markdown"]
    assert data["info"]["type"] == "word"
    assert data["info"]["headings"] >= 1
    assert data["info"]["tables"] >= 1


def test_convert_direct_xlsx():
    xlsx_bytes = make_xlsx_bytes()
    r = client.post(
        "/api/convert-direct",
        files={"file": ("data.xlsx", xlsx_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
    )
    assert r.status_code == 200
    data = r.json()
    assert data["filename"] == "data.xlsx"
    assert "## Sheet: Sales" in data["markdown"]
    assert "Name" in data["markdown"]
    assert "Alice" in data["markdown"]
    assert data["info"]["type"] == "excel"
    assert len(data["info"]["sheets"]) == 1
    assert data["info"]["sheets"][0]["name"] == "Sales"
    assert data["info"]["sheets"][0]["rows"] == 3
