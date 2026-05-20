#!/usr/bin/env python3
"""
convert.py — Batch document → Markdown converter.

Supported input formats:
  • PDF            — OCR via Ollama GLM-OCR (always)
  • DOCX / DOC     — direct conversion (python-docx) or OCR via LibreOffice+Ollama
  • XLSX / XLS     — direct conversion (openpyxl) or OCR via LibreOffice+Ollama

Conversion modes (--mode):
  auto    Use direct conversion for DOCX/XLSX; OCR for PDF/DOC/XLS  [default]
  direct  Force direct conversion for DOCX/XLSX (errors on other formats)
  ocr     Force OCR pipeline for every file (requires LibreOffice for Office docs)

Usage:
    python convert.py [options]

Options:
    -i, --input   DIR      Input folder  (default: ../raw)
    -o, --output  DIR      Output folder (default: ../output)
    -f, --force            Re-process files that already have an output
    --mode        MODE     Conversion mode: auto | direct | ocr  (default: auto)
    --dpi         INT      Render resolution in DPI for OCR (default: 200)
    --model       STR      Ollama model name (default: glm-ocr:latest)
    --ollama-url  URL      Ollama endpoint (default: http://localhost:11434/api/generate)
    --timeout     SEC      Per-page OCR timeout in seconds (default: 180)
"""

import argparse
import base64
import io
import os
import subprocess
import sys
import tempfile
import time
from pathlib import Path

import fitz  # PyMuPDF
import httpx

# ── Defaults ──────────────────────────────────────────────────────────────────

SCRIPTS_DIR = Path(__file__).resolve().parent
ROOT_DIR    = SCRIPTS_DIR.parent

DEFAULT_INPUT      = ROOT_DIR / "raw"
DEFAULT_OUTPUT     = ROOT_DIR / "output"
DEFAULT_MODEL      = "glm-ocr:latest"
DEFAULT_OLLAMA_URL = "http://localhost:11434/api/generate"
DEFAULT_DPI        = 200
DEFAULT_TIMEOUT    = 180

# Extensions handled by each strategy
DIRECT_EXTS = {".docx", ".xlsx"}
OCR_ONLY_EXTS = {".doc", ".xls", ".pdf"}
ALL_EXTS = DIRECT_EXTS | OCR_ONLY_EXTS

OCR_PROMPT = (
    "Convert the content of this document page to clean, well-structured Markdown. "
    "Preserve all text exactly. Format tables using Markdown table syntax (| col | col |). "
    "Use appropriate heading levels (#, ##, ###). Format lists with - or 1. "
    "Wrap code snippets in backticks. For figures or diagrams write [Figure: brief description]. "
    "Output only the Markdown content, no preamble or explanation."
)

# ── Colours (graceful fallback on non-ANSI terminals) ────────────────────────

def _c(code: str, text: str) -> str:
    return f"\033[{code}m{text}\033[0m" if sys.stdout.isatty() else text

def green(t):  return _c("32", t)
def yellow(t): return _c("33", t)
def red(t):    return _c("31", t)
def bold(t):   return _c("1",  t)
def dim(t):    return _c("2",  t)

# ── Direct converters (DOCX / XLSX) ──────────────────────────────────────────

def _runs_to_md(runs) -> str:
    md = ""
    for run in runs:
        t = run.text
        if not t:
            continue
        if run.bold and run.italic:
            t = f"***{t}***"
        elif run.bold:
            t = f"**{t}**"
        elif run.italic:
            t = f"*{t}*"
        md += t
    return md


def _table_to_md(table) -> str:
    rows = table.rows
    if not rows:
        return ""
    def _cell(c):
        return c.text.strip().replace("\n", " ").replace("|", "\\|")
    header = "| " + " | ".join(_cell(c) for c in rows[0].cells) + " |"
    sep    = "| " + " | ".join("---" for _ in rows[0].cells) + " |"
    data   = [
        "| " + " | ".join(_cell(c) for c in row.cells) + " |"
        for row in rows[1:]
    ]
    return "\n".join([header, sep] + data)


HEADING_STYLES = {
    "Heading 1": "#",  "Heading 2": "##",  "Heading 3": "###",
    "Heading 4": "####", "Heading 5": "#####", "Heading 6": "######",
    "Title": "#",      "Subtitle": "##",
}


def docx_to_markdown(path: Path) -> tuple[str, dict]:
    from docx import Document
    content = path.read_bytes()
    doc = Document(io.BytesIO(content))
    parts = []
    info = {"type": "word", "paragraphs": 0, "headings": 0, "tables": 0}

    para_idx = table_idx = 0
    for child in doc.element.body:
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag == "p":
            if para_idx < len(doc.paragraphs):
                para = doc.paragraphs[para_idx]
                para_idx += 1
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else ""
                if style_name in HEADING_STYLES:
                    parts.append(f"{HEADING_STYLES[style_name]} {text}")
                    info["headings"] += 1
                else:
                    md_text = _runs_to_md(para.runs)
                    if md_text.strip():
                        parts.append(md_text)
                        info["paragraphs"] += 1
        elif tag == "tbl":
            if table_idx < len(doc.tables):
                tbl = doc.tables[table_idx]
                table_idx += 1
                md_table = _table_to_md(tbl)
                if md_table:
                    parts.append(md_table)
                    info["tables"] += 1

    return "\n\n".join(parts), info


def xlsx_to_markdown(path: Path) -> tuple[str, dict]:
    import openpyxl
    content = path.read_bytes()
    wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
    parts = []
    info = {"type": "excel", "sheets": []}

    def _cell_str(c):
        return (str(c) if c is not None else "").replace("\n", " ").replace("|", "\\|")

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        non_empty = [r for r in rows if any(c is not None for c in r)]
        info["sheets"].append({
            "name": sheet_name,
            "rows": len(non_empty),
            "cols": len(non_empty[0]) if non_empty else 0,
        })
        if not non_empty:
            continue
        parts.append(f"## Sheet: {sheet_name}")
        header = non_empty[0]
        parts.append("| " + " | ".join(_cell_str(c) for c in header) + " |")
        parts.append("| " + " | ".join("---" for _ in header) + " |")
        for row in non_empty[1:]:
            parts.append("| " + " | ".join(_cell_str(c) for c in row) + " |")

    return "\n\n".join(parts), info

# ── OCR helpers ───────────────────────────────────────────────────────────────

def _office_to_pdf_bytes(path: Path) -> bytes:
    """Convert a DOC/DOCX/XLS/XLSX to PDF bytes via LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdir:
        input_path = os.path.join(tmpdir, path.name)
        with open(input_path, "wb") as f:
            f.write(path.read_bytes())
        try:
            result = subprocess.run(
                ["libreoffice", "--headless", "--convert-to", "pdf",
                 "--outdir", tmpdir, input_path],
                capture_output=True, timeout=120,
            )
        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice timed out after 120 seconds")
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice failed: {result.stderr.decode(errors='replace')}"
            )
        pdf_path = os.path.join(tmpdir, path.stem + ".pdf")
        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice did not produce a PDF file")
        with open(pdf_path, "rb") as f:
            return f.read()


def render_pages(pdf_bytes: bytes, dpi: int) -> list[tuple[int, str]]:
    """Return [(page_number, base64_jpeg), ...] for every page."""
    mat = fitz.Matrix(dpi / 72, dpi / 72)
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pages = []
    try:
        for i in range(len(doc)):
            pix  = doc[i].get_pixmap(matrix=mat)
            jpeg = pix.tobytes("jpeg", jpg_quality=90)
            pages.append((i + 1, base64.b64encode(jpeg).decode()))
    finally:
        doc.close()
    return pages


def ocr_page(b64_image: str, *, model: str, ollama_url: str, timeout: int) -> str:
    payload = {
        "model":  model,
        "prompt": OCR_PROMPT,
        "images": [b64_image],
        "stream": False,
    }
    with httpx.Client(timeout=timeout) as client:
        resp = client.post(ollama_url, json=payload)
        resp.raise_for_status()
        return resp.json().get("response", "").strip()


def check_ollama(ollama_url: str, model: str) -> None:
    base = ollama_url.rstrip("/").removesuffix("/api/generate")
    try:
        r = httpx.get(f"{base}/api/tags", timeout=10)
        r.raise_for_status()
        names = [m["name"] for m in r.json().get("models", [])]
        if not any(n.startswith(model.split(":")[0]) for n in names):
            print(yellow(f"  Warning: model '{model}' not found in Ollama. "
                         f"Available: {', '.join(names) or 'none'}"))
    except httpx.ConnectError:
        print(red(f"  Error: cannot reach Ollama at {ollama_url}"))
        print(red("  Make sure 'ollama serve' is running."))
        sys.exit(1)

# ── Per-file conversion ───────────────────────────────────────────────────────

def _effective_mode(path: Path, mode: str) -> str:
    """Return the actual strategy ('direct' or 'ocr') for a given file."""
    ext = path.suffix.lower()
    if mode == "direct":
        if ext not in DIRECT_EXTS:
            raise ValueError(
                f"--mode direct is not supported for {ext} files. "
                "Use 'auto' or 'ocr'."
            )
        return "direct"
    if mode == "ocr":
        return "ocr"
    # auto
    return "direct" if ext in DIRECT_EXTS else "ocr"


def convert_file(
    path: Path,
    output_dir: Path,
    *,
    mode: str,
    dpi: int,
    model: str,
    ollama_url: str,
    timeout: int,
) -> dict:
    t0  = time.time()
    ext = path.suffix.lower()
    strategy = _effective_mode(path, mode)

    # ── Direct conversion ──
    if strategy == "direct":
        print(f"  Converting directly…", end=" ", flush=True)
        if ext == ".docx":
            markdown, info = docx_to_markdown(path)
            label = f"Word ({info['paragraphs']} paragraphs, {info['headings']} headings, {info['tables']} tables)"
        elif ext == ".xlsx":
            markdown, info = xlsx_to_markdown(path)
            sheets = ", ".join(f"{s['name']} ({s['rows']}r×{s['cols']}c)" for s in info["sheets"])
            label = f"Excel ({len(info['sheets'])} sheet(s): {sheets})"
        else:
            raise ValueError(f"Direct conversion not supported for {ext}")

        output_path = output_dir / (path.stem + ".md")
        output_path.write_text(markdown, encoding="utf-8")
        elapsed = time.time() - t0
        print(green("done") + dim(f"  ({elapsed:.1f}s)"))
        print(dim(f"  {label}"))
        return {
            "strategy": "direct",
            "pages": 1,
            "failed": 0,
            "failed_pages": [],
            "duration": elapsed,
            "output": output_path,
        }

    # ── OCR pipeline ──
    if ext == ".pdf":
        print(f"  Reading PDF…", end=" ", flush=True)
        pdf_bytes = path.read_bytes()
    else:
        print(f"  Converting {ext} → PDF via LibreOffice…", end=" ", flush=True)
        pdf_bytes = _office_to_pdf_bytes(path)

    pages = render_pages(pdf_bytes, dpi)
    print(f"{len(pages)} page(s) ready.")

    sections: list[str] = []
    failed_pages: list[int] = []

    for page_num, b64 in pages:
        prefix = f"  Page {page_num:>3}/{len(pages)}"
        print(f"{prefix}  OCR… ", end="", flush=True)
        pt0 = time.time()
        try:
            md = ocr_page(b64, model=model, ollama_url=ollama_url, timeout=timeout)
            print(green("done") + dim(f"  ({time.time() - pt0:.1f}s)"))
            sections.append(f"<!-- page {page_num} -->\n\n{md}")
        except httpx.TimeoutException:
            print(yellow("timeout — skipped"))
            failed_pages.append(page_num)
            sections.append(f"<!-- page {page_num} -->\n\n*(OCR timed out)*")
        except Exception as exc:
            print(red(f"error — {exc}"))
            failed_pages.append(page_num)
            sections.append(f"<!-- page {page_num} -->\n\n*(OCR failed: {exc})*")

    markdown = "\n\n---\n\n".join(sections)
    output_path = output_dir / (path.stem + ".md")
    output_path.write_text(markdown, encoding="utf-8")

    return {
        "strategy": "ocr",
        "pages": len(pages),
        "failed": len(failed_pages),
        "failed_pages": failed_pages,
        "duration": time.time() - t0,
        "output": output_path,
    }

# ── CLI ───────────────────────────────────────────────────────────────────────

def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(
        description=(
            "Batch-convert PDF, DOCX, DOC, XLSX, XLS files to Markdown. "
            "Uses direct conversion for DOCX/XLSX (auto mode) and "
            "GLM-OCR via Ollama for PDFs and legacy formats."
        ),
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )
    p.add_argument("-i", "--input",      default=str(DEFAULT_INPUT),      metavar="DIR")
    p.add_argument("-o", "--output",     default=str(DEFAULT_OUTPUT),     metavar="DIR")
    p.add_argument("-f", "--force",      action="store_true",
                   help="Re-process files that already have an output")
    p.add_argument("--mode",            default="auto",
                   choices=["auto", "direct", "ocr"],
                   help="auto=smart, direct=python-docx/openpyxl, ocr=Ollama pipeline")
    p.add_argument("--dpi",             default=DEFAULT_DPI,   type=int,  metavar="INT")
    p.add_argument("--model",           default=DEFAULT_MODEL,            metavar="STR")
    p.add_argument("--ollama-url",      default=DEFAULT_OLLAMA_URL,       metavar="URL")
    p.add_argument("--timeout",         default=DEFAULT_TIMEOUT, type=int, metavar="SEC")
    return p.parse_args()


def main() -> None:
    args = parse_args()

    raw_dir    = Path(args.input).expanduser().resolve()
    output_dir = Path(args.output).expanduser().resolve()

    if not raw_dir.is_dir():
        print(red(f"Input folder not found: {raw_dir}"))
        sys.exit(1)

    output_dir.mkdir(parents=True, exist_ok=True)

    docs = sorted(f for f in raw_dir.iterdir() if f.suffix.lower() in ALL_EXTS)
    if not docs:
        print(yellow(f"No supported files found in {raw_dir}"))
        print(dim(f"  Supported: {', '.join(sorted(ALL_EXTS))}"))
        sys.exit(0)

    print(bold(f"\ndoc2mrk — batch converter"))
    print(dim(f"  input  : {raw_dir}"))
    print(dim(f"  output : {output_dir}"))
    print(dim(f"  mode   : {args.mode}  |  DPI: {args.dpi}  |  timeout: {args.timeout}s"))

    needs_ocr = args.mode in ("ocr", "auto") and any(
        f.suffix.lower() in OCR_ONLY_EXTS or
        (args.mode == "ocr" and f.suffix.lower() in DIRECT_EXTS)
        for f in docs
    )
    if needs_ocr:
        check_ollama(args.ollama_url, args.model)

    # ── Filter already-done files ──
    to_process, skipped = [], []
    for doc in docs:
        out = output_dir / (doc.stem + ".md")
        (skipped if out.exists() and not args.force else to_process).append(doc)

    if skipped:
        print(dim(f"\n  Skipping {len(skipped)} already-converted file(s) "
                  f"(use --force to re-process):"))
        for f in skipped:
            print(dim(f"    • {f.name}"))

    if not to_process:
        print(green("\n  All files already converted. Nothing to do."))
        sys.exit(0)

    print(f"\n  Converting {bold(str(len(to_process)))} file(s)…")

    total_t0 = time.time()
    results  = {}

    for idx, doc in enumerate(to_process, 1):
        ext = doc.suffix.lower()
        strategy = _effective_mode(doc, args.mode)
        tag = dim(f"[{strategy}]")
        print(f"\n{'─' * 60}")
        print(f"  [{idx}/{len(to_process)}] {bold(doc.name)}  {tag}")

        try:
            r = convert_file(
                doc, output_dir,
                mode=args.mode,
                dpi=args.dpi,
                model=args.model,
                ollama_url=args.ollama_url,
                timeout=args.timeout,
            )
            results[doc.name] = r
            if r["strategy"] == "direct":
                status = green("✓ OK")
            elif r["failed"] == 0:
                status = green("✓ OK")
            else:
                status = yellow(f"✓ done  ({r['failed']} page(s) failed: {r['failed_pages']})")
            print(f"\n  {status}  →  {dim(str(r['output']))}")
            print(dim(f"  {r['pages']} page(s) in {r['duration']:.1f}s"))
        except Exception as exc:
            print(red(f"\n  ✗ Failed: {exc}"))
            results[doc.name] = {"error": str(exc)}

    # ── Summary ──
    total_elapsed = time.time() - total_t0
    ok   = sum(1 for r in results.values() if "error" not in r and r.get("failed", 0) == 0)
    warn = sum(1 for r in results.values() if "error" not in r and r.get("failed", 0) > 0)
    fail = sum(1 for r in results.values() if "error" in r)

    print(f"\n{'═' * 60}")
    print(bold("  Summary"))
    print(f"  {green(f'{ok} converted')}  "
          f"{yellow(f'{warn} with warnings') if warn else ''}  "
          f"{red(f'{fail} failed') if fail else ''}".rstrip())
    print(dim(f"  Total time: {total_elapsed:.1f}s"))
    print(f"  Output in: {output_dir}\n")

    sys.exit(1 if fail else 0)


if __name__ == "__main__":
    main()
